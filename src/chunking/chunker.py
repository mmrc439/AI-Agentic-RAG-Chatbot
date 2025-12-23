from src.data_ingest.preprocess import (
    _normalize_heading_title,
    autodetect_page_numbering,
    extract_page_numbers_from_headers_footers,
    first_body_index_from_labels,
    make_toc_whitelist,
)
from src.data_ingest.preprocess import (
    sha256_file,
    find_toc_pages,
    gpt_layout_probe,
    render_pages_to_images,
    bands_for_page,
    select_optimal_probe_pages,
    _normalize_heading_number,
    now_utc_iso,
    is_pure_number_token,
    split_text_tokens,  # retained for parity; iterator defined below
    stable_hash,
    validate_heading_patterns,
    _augment_heading_patterns_to_toc_depth,
    cluster_font_sizes,
    clean_text,
    is_toc_like_line,
    is_leaderless_toc_line,
    numeric_vector_from_heading,
    enc,                 # tokenizer to compute token offsets for deterministic IDs
    heading_threshold,   # explicitly imported; used in heading font guard
    _tok_len,            # shared token-length helper (safe wrapper over tokenizer)
    effective_depth_from_vector,
    structurally_heading_like,
    render_bbox_image,
    find_pdf_caption_near,
)
from src.configs.config import DEBUG
from src.configs.helper import *  # brings CAPTION_PATTERNS, MAX_TOKENS, OVERLAP, RUN_GPT_* flags, GPT_MODEL, STORE_DIR, etc.
from chromadb import Settings
from langchain_core.documents import Document

import fitz
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any, Iterator, Tuple as TypingTuple
import hashlib, json, os, re, io, base64
from PIL import Image
from openai import OpenAI  # used for GPT-Vision-style image description


# ──────────────────────────────────────────────────────────────────────────────
# Backwards-compat constant aliases (so other modules importing from chunker
# still see the same names as before)
# ──────────────────────────────────────────────────────────────────────────────
KEEP_TOP_K         = KEEP_TOP_K
SIM_THRESHOLD      = SIM_THRESHOLD
CITE_SIM_THRESHOLD = CITE_SIM_THRESHOLD
SCOPED_MAX         = SCOPED_MAX
GLOBAL_K           = GLOBAL_K
MAX_TOKENS         = MAX_TOKENS
OVERLAP            = OVERLAP
REQ_BULLET_LIMIT   = REQ_BULLET_LIMIT
LLM_PROVIDER       = LLM_PROVIDER
RUN_GPT_LAYOUT_PROBE   = RUN_GPT_LAYOUT_PROBE
RUN_GPT_IMAGE_DESCRIBE = RUN_GPT_IMAGE_DESCRIBE
RUN_GPT_DOC_SUMMARY    = RUN_GPT_DOC_SUMMARY
GPT_MODEL          = GPT_MODEL
STORE_DIR          = STORE_DIR


# Optional: same wrapper class you had before, kept for compatibility
if LLM_PROVIDER == "OPENAI":
    class OpenAIWrapper:
        def __init__(self, model_name: str):
            self.client = OpenAI()
            self.model_name = model_name

        def invoke(self, prompt, **kwargs) -> str:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=prompt,
                temperature=kwargs.get("temperature", 0.2),
                max_tokens=kwargs.get("max_tokens", 512),
                top_p=kwargs.get("top_p", 0.9),
            )
            # Keeping your original access pattern
            return response.choices[0].message["content"]


# ──────────────────────────────────────────────────────────────────────────────
# OpenAI client for image description (as in old chunker)
# ──────────────────────────────────────────────────────────────────────────────
try:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    print("[INFO] OpenAI client initialized for image description")
except Exception as e:
    client = None
    print(f"[WARN] OpenAI client initialization failed: {e}")


# ──────────────────────────────────────────────────────────────────────────────
# Constants / regexes
# ──────────────────────────────────────────────────────────────────────────────
CAPTION_RE_LIST = [re.compile(p, re.I) for p in CAPTION_PATTERNS]

# Lightweight quality gate knobs (safe defaults)
MIN_CHARS_PER_CHUNK  = 20
MIN_TOKENS_PER_CHUNK = 8

# Header/footer de-dup knobs
HF_CANDIDATE_MAX_LEN    = 80   # characters
HF_MIN_PAGES_FOR_REPEAT = 3    # once a line shows up on ≥ this many pages, treat as repeatable boilerplate


# ──────────────────────────────────────────────────────────────────────────────
# Deterministic token-window iterator (returns (text, start_token_offset))
# ──────────────────────────────────────────────────────────────────────────────
def iter_token_windows(text: str, max_tokens: int, overlap: int) -> Iterator[Tuple[str, int]]:
    """
    Deterministic windowing with token start offsets so chunk_id can be stable.
    Uses exact token offsets; windows are contiguous with step=(max_tokens-overlap).
    """
    if not (text or "").strip():
        return
    toks = enc.encode(text)
    step = max(1, max_tokens - overlap)
    i = 0
    while i < len(toks):
        j = min(len(toks), i + max_tokens)
        yield enc.decode(toks[i:j]), i
        if j == len(toks):
            break
        i += step


# ──────────────────────────────────────────────────────────────────────────────
# Header/Footer repeating-line filter
# ──────────────────────────────────────────────────────────────────────────────
class RepeatingLineFilter:
    """
    Tracks short lines seen across pages. When a line appears on >= HF_MIN_PAGES_FOR_REPEAT
    different pages, we mark it as a repeating header/footer and filter it subsequently.
    """
    def __init__(self):
        self.line_pages: Dict[str, set] = {}      # line -> {page_idx,...}
        self.blocklist: set = set()               # lines confirmed as repeating
        self.dropped_count = 0

    def consider(self, line: str, page_idx: int):
        s = (line or "").strip()
        # only short lines that look header/footer-ish
        if not s or len(s) > HF_CANDIDATE_MAX_LEN:
            return
        st = s.lower()
        pages = self.line_pages.setdefault(st, set())
        pages.add(page_idx)
        if (st not in self.blocklist) and (len(pages) >= HF_MIN_PAGES_FOR_REPEAT):
            self.blocklist.add(st)

    def should_drop(self, line: str) -> bool:
        st = (line or "").strip().lower()
        if st in self.blocklist:
            self.dropped_count += 1
            return True
        return False


# ──────────────────────────────────────────────────────────────────────────────
# Small helpers
# ──────────────────────────────────────────────────────────────────────────────
def _is_probably_code(text: str) -> bool:
    if not text:
        return False
    SYMBOLS_FOR_CODE = r"[{}\[\]();=:<>/*#`~|\\]"
    CODE_SYMBOL_RATIO = 0.10
    CODE_INDENT_SHARE = 0.30
    symbol_ratio = len(re.findall(SYMBOLS_FOR_CODE, text)) / max(1, len(text))
    lines = [l for l in text.splitlines() if l.strip()]
    indented = sum(1 for l in lines if re.match(r"^(\t| {2,})", l)) if lines else 0
    indent_share = (indented / len(lines)) if lines else 0.0
    return symbol_ratio >= CODE_SYMBOL_RATIO or indent_share >= CODE_INDENT_SHARE


def _describe_image_with_gpt(blob: bytes, mime_hint: Optional[str] = None) -> Optional[dict]:
    """
    Old behavior: returns dict with keys depending on modality:
      image: {"modality":"image","caption":str,"long_description":str}
      code : {"modality":"code","code":str,"language":str?,"caption":str?}
      table: {"modality":"table","columns":[...],"rows":[[...]],"caption":str?}
    """
    if not RUN_GPT_IMAGE_DESCRIBE or client is None:
        return None
    try:
        # Normalize to PNG for reliability
        try:
            im = Image.open(io.BytesIO(blob)).convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="PNG")
            png_bytes = buf.getvalue()
            mime = "image/png"
        except Exception:
            png_bytes = blob
            mime = mime_hint or "image/png"

        data_url = "data:%s;base64,%s" % (mime, base64.b64encode(png_bytes).decode("utf-8"))

        system = (
            "You convert images into structured JSON. Pick ONE modality: image, code, or table.\n"
            "- If code: OCR exact code (no paraphrase), infer language if obvious.\n"
            "- If table: extract headers (columns) and rows faithfully (cap at ~20 rows).\n"
            "- If image: produce BOTH a short caption (<=120 chars) AND a long_description (concise but detailed: "
            "chart/diagram type, axes, units, key values if legible, visible trends, relationships, anomalies, "
            "and any clearly readable text). Do NOT invent numbers; only include values that are visible.\n"
            "Return ONLY valid JSON."
        )
        user = [
            {
                "type": "text",
                "text": (
                    "{\n"
                    '  "modality": "image"|"code"|"table",\n'
                    '  "caption": string (short; for image/code/table),\n'
                    '  "long_description": string (for modality="image"),\n'
                    '  "language": string (for code),\n'
                    '  "code": string (for code),\n'
                    '  "columns": string[] (for table),\n'
                    '  "rows": string[][] (for table)\n'
                    "}\n"
                    "Constraints:\n"
                    "- caption <= 120 chars\n"
                    "- long_description ~150-400 words; no hallucinated numbers; summarize trends precisely\n"
                    "- rows <= 20, columns <= 30\n"
                ),
            },
            {"type": "image_url", "image_url": {"url": data_url}},
        ]
        resp = client.chat.completions.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=0,
            response_format={"type": "json_object"},
        )
        payload = json.loads(resp.choices[0].message.content)

        mod = (payload.get("modality") or "image").lower().strip()
        if mod not in ("image", "code", "table"):
            mod = "image"

        # enforce basic caps to avoid giant fields
        def _clip(s, n):
            return (s[:n].rstrip() + "…" if isinstance(s, str) and len(s) > n else s)

        out = {"modality": mod}
        if mod == "code":
            out["code"] = payload.get("code", "") or ""
            out["language"] = payload.get("language")
            out["caption"] = _clip(payload.get("caption") or "", 200)
        elif mod == "table":
            out["columns"] = [str(c) for c in (payload.get("columns") or [])][:30]
            out["rows"] = [[str(c) for c in r] for r in (payload.get("rows") or [])][:20]
            out["caption"] = _clip(payload.get("caption") or "", 200)
        else:
            out["caption"] = _clip(payload.get("caption") or "", 200)
            out["long_description"] = _clip(payload.get("long_description") or "", 8000)
        return out
    except Exception as e:
        dbg(f"[DBG] GPT-vision failed: {e}")
        return None


def _demote_loose_L1_if_noisy(pdf, compiled_by_level, rules, page_numbering_info, body_start) -> Dict[int, List[re.Pattern]]:
    if 1 not in compiled_by_level:
        return compiled_by_level
    loose_patterns = [p for p in compiled_by_level.get(1, []) if p.pattern == r"^\d+\s+"]
    has_strict = (2 in compiled_by_level) or (3 in compiled_by_level)
    if not loose_patterns or not has_strict:
        return compiled_by_level

    sample_pages = [i for i in [body_start, body_start + 1] if 0 <= i < pdf.page_count]
    if not sample_pages:
        return compiled_by_level

    header_present, footer_present = True, True
    l1_hits, l1_noise = 0, 0
    for i in sample_pages:
        pg = pdf.load_page(i)
        _, _, hb, fb, content_region = bands_for_page(pg, rules)
        blocks = pg.get_text("dict").get("blocks", [])
        for blk in blocks:
            if blk.get("type") != 0 or not blk.get("bbox"):
                continue
            x0, y0, x1, y1 = blk["bbox"]
            if x1 < content_region[0] or x0 > content_region[2]:
                continue
            if y1 < content_region[1] or y0 > content_region[3]:
                continue
            ym = (y0 + y1) / 2.0
            if header_present and hb[0] <= ym <= hb[1]:
                continue
            if footer_present and fb[0] <= ym <= fb[1]:
                continue
            for ln in blk.get("lines", []) or []:
                spans = ln.get("spans", []) or []
                text = clean_text("".join(sp.get("text", "") for sp in spans))
                if not text:
                    continue
                if any(p.match(text) for p in loose_patterns):
                    l1_hits += 1
                    words = text.split()
                    long_line = len(text) > 120 or len(words) > 18
                    sentencey = bool(re.search(r"[a-z].*[\.!?]$", text))
                    year_vec = numeric_vector_from_heading(text)
                    year_like = bool(year_vec and len(year_vec) == 1 and 1900 <= year_vec[0] <= 2100)
                    if long_line or sentencey or year_like:
                        l1_noise += 1

    if l1_hits >= 4 and (l1_noise / max(1, l1_hits)) >= 0.5:
        compiled_by_level[1] = [p for p in compiled_by_level[1] if p.pattern != r"^\d+\s+"]
        if not compiled_by_level[1]:
            del compiled_by_level[1]
        dbg(r"[DBG] Dropping loose L1 pattern ^\d+\s+ (too permissive)")
    return compiled_by_level


def _page_kind_with_override(page_idx, body_start, hf_label, is_toc_candidate=False):
    if is_toc_candidate:
        return "front_matter"
    if (
        re.match(r"^\d+$", str(hf_label or ""))
        and int(hf_label) >= 1
        and page_idx >= max(0, (body_start or 0) - 1)
    ):
        return "body"
    return "front_matter" if page_idx < (body_start or 0) else "body"


# ──────────────────────────────────────────────────────────────────────────────
# Stable doc_id helper
# ──────────────────────────────────────────────────────────────────────────────
def _make_doc_id(path: Path, file_hash: str) -> str:
    stem = path.stem
    return f"{stem}::{file_hash[:6]}" if file_hash else stem


# ──────────────────────────────────────────────────────────────────────────────
# Buffered text → chunks with deterministic IDs + quality gates
# ──────────────────────────────────────────────────────────────────────────────
def _flush_text_buffer(
    chunks: List[dict],
    current_text_ref: Dict[str, str],
    section_info_fn,
    *,
    file_hash: str,
    page_num: int,
    page_label: str,
    page_type: str,
    page_idx: int,
    doc_name: str,
    ingested_at: str,
    doc_id: str,
    dedup_set: set,
    drop_counters: Dict[str, int],
):
    txt = current_text_ref["buf"]
    if not (txt or "").strip():
        return

    si = section_info_fn()
    modality = "code" if _is_probably_code(txt) else "text"

    # Iterate token windows with offsets for deterministic IDs
    for piece, start_off in iter_token_windows(txt, MAX_TOKENS, OVERLAP):
        # Quality gates
        if len(piece) < MIN_CHARS_PER_CHUNK:
            drop_counters["too_short_chars"] += 1
            continue
        if _tok_len(piece) < MIN_TOKENS_PER_CHUNK:
            drop_counters["too_short_tokens"] += 1
            continue

        # Exact intra-doc dedup
        ph = hashlib.sha1(piece.encode("utf-8")).hexdigest()
        if ph in dedup_set:
            drop_counters["exact_dup"] += 1
            continue
        dedup_set.add(ph)

        # Deterministic ID (doc_id + page label + start token offset) + content salt for uniqueness
        ph_short = ph[:8]
        cid = f"{doc_id}::{page_label}::t{start_off}:{ph_short}"

        chunks.append(
            {
                **si,
                "doc_name": doc_name,
                "doc_type": "pdf",
                "file_hash": file_hash,
                "doc_id": doc_id,
                "page_index": page_num,
                "page_label": page_label,
                "page_type": page_type,
                "modality": modality,
                "kind": "code" if modality == "code" else "text",
                "section_title": si.get("section"),
                "heading_path": si.get("section_path"),
                "content": piece,
                "caption_text": None,
                "bounding_box": None,
                "ingested_at": ingested_at,
                "chunk_id": cid,
                "chunk_hash": cid,  # keep same for downstream
            }
        )

    current_text_ref["buf"] = ""


# ──────────────────────────────────────────────────────────────────────────────
# Main: PDF → chunks (text + tables + images/diagrams via GPT)
# ──────────────────────────────────────────────────────────────────────────────
def parse_pdf_to_chunks(path: str) -> List[dict]:
    path = Path(path)
    pdf = fitz.open(str(path))
    try:
        file_hash = sha256_file(path)
        doc_name = path.name
        doc_id = _make_doc_id(path, file_hash)

        dbg(f"[DBG] Processing PDF: {doc_name}")

        # Header/footer repeating-line filter
        hf_filter = RepeatingLineFilter()

        # TOC discovery
        toc_pages, toc_meta = find_toc_pages(pdf)
        if not toc_pages:
            for i in range(min(pdf.page_count, 5)):
                text = (pdf.load_page(i).get_text("text") or "").strip()
                first = (text.splitlines() or [""])[0][:80]
                if re.search(r"^\s*Index\s*$", first, re.I):
                    toc_pages = [i]
                    toc_meta = {"source": "fallback_index_title"}
                    break

        # Pre-probe & layout probe
        preprobe = []
        if toc_pages:
            pg0_text = (pdf.load_page(toc_pages[0]).get_text("text") or "").strip()
            preprobe.append(toc_pages[0] if len(pg0_text) > 10 else (1 if pdf.page_count > 1 else 0))
        probe_base = [p for p in [0, min(1, pdf.page_count - 1)] if p not in preprobe]
        probe_pages = list(dict.fromkeys((preprobe + probe_base)[:3])) or [0, 1]
        rules = gpt_layout_probe(render_pages_to_images(pdf, probe_pages))

        # Page numbering
        page_numbering_info = rules.get("page_numbering") or {}
        if not page_numbering_info or page_numbering_info.get("location", "unknown") == "unknown":
            page_numbering_info = autodetect_page_numbering(pdf)

        # Body start via header/footer labels
        labels_by_page: Dict[int, str] = {}
        body_start: Optional[int] = None
        for i in range(pdf.page_count):
            pg = pdf.load_page(i)
            _, _, hb, fb, _ = bands_for_page(pg, rules)
            lab = extract_page_numbers_from_headers_footers(pg, hb, fb, page_numbering_info)
            if lab:
                labels_by_page[i] = lab
                if lab == "1" and body_start is None:
                    body_start = i
        if body_start is None:
            tmp = first_body_index_from_labels(pdf)
            if tmp is not None:
                body_start = tmp
        if body_start is None:
            body_start = 0

        # Reconcile TOC pages
        orig_toc_pages = toc_pages[:]
        toc_pages = [p for p in toc_pages if p < body_start]
        skip_pages = set(toc_pages)

        # TOC whitelist & hierarchy maps
        toc_whitelist = make_toc_whitelist(pdf, toc_pages)
        by_num: Dict[str, dict] = {}
        by_title_at_level: Dict[Tuple[int, str], dict] = {}
        if toc_whitelist.get("has_toc"):
            idx = toc_whitelist.get("idx") or []
            stack_path: List[dict] = []
            for e in idx:
                lvl = int(e.get("level") or 1)
                num = _normalize_heading_number(e.get("title", "")) or e.get("num")
                title_raw = (e.get("title") or "").strip()
                title_norm = _normalize_heading_title(title_raw) or title_raw
                while len(stack_path) >= lvl:
                    stack_path.pop()
                stack_path.append({"num": num, "title": title_raw})
                path_titles = [d["title"] for d in stack_path]
                path_nums = [d["num"] for d in stack_path]
                node = {
                    "level": lvl,
                    "num": num,
                    "title": title_raw,
                    "path_titles": path_titles[:],
                    "path_nums": path_nums[:],
                }

                if num:
                    by_num.setdefault(num, node)
                key = (lvl, _normalize_heading_title(title_raw))
                by_title_at_level.setdefault(key, []).append(node)

        # Better probe on hierarchy pages
        probe_pages = select_optimal_probe_pages(pdf, toc_pages, body_start)
        rules = gpt_layout_probe(render_pages_to_images(pdf, probe_pages))
        page_numbering_info = rules.get("page_numbering") or page_numbering_info

        # Heading patterns
        compiled_by_level = validate_heading_patterns(rules)
        if toc_whitelist["has_toc"]:
            compiled_by_level = _augment_heading_patterns_to_toc_depth(
                compiled_by_level, toc_whitelist["max_depth"]
            )
        compiled_by_level = _demote_loose_L1_if_noisy(
            pdf, compiled_by_level, rules, page_numbering_info, body_start
        )

        font_stats = cluster_font_sizes(pdf)
        heading_thresh = heading_threshold(
            font_stats, rules["heading_detection"].get("font_size_hint_ratio")
        )

        # Output accumulators
        chunks: List[dict] = []
        section_stack: List[dict] = []
        current_text = {"buf": ""}
        ingested_at = now_utc_iso()
        doc_dedup_set: set = set()
        drop_counters = {"too_short_chars": 0, "too_short_tokens": 0, "exact_dup": 0}

        dotted_drops = 0
        leaderless_drops = 0
        fused_count = 0

        # section info
        def section_info():
            if not section_stack:
                return {
                    "section": "Document",
                    "section_level": 0,
                    "section_path": "Document",
                    "section_type": "content",
                }

            def _piece(s: dict) -> str:
                num = (s.get("num") or "").rstrip(".").strip()
                title = (s.get("title") or "").strip()
                if not num:
                    return title
                if (
                    not title
                    or title == num
                    or title.lower().startswith((num + " ").lower())
                ):
                    return title or num
                return f"{num} {title}"

            path_parts: List[str] = []
            acc: List[dict] = []
            for s in section_stack:
                while acc and acc[-1]["level"] >= s["level"]:
                    acc.pop()
                    if path_parts:
                        path_parts.pop()
                acc.append(s)
                path_parts.append(_piece(s))
            return {
                "section": (section_stack[-1].get("title") or "").strip(),
                "section_level": section_stack[-1]["level"],
                "section_path": " / ".join(p for p in path_parts if p),
                "section_type": section_stack[-1].get("section_type", "content"),
            }

        last_ctx = {"page_num": None, "page_label": None, "page_type": None, "page_idx": None}

        # Local helpers
        def append_body_text(s: str):
            if not s:
                return
            # Header/Footer last-pass: strip repeating short lines
            if hf_filter.should_drop(s):
                return
            hf_filter.consider(s, last_ctx["page_idx"])

            buf = current_text["buf"]
            candidate = (buf + "\n" + s).strip() if buf else s
            # If we are near the cap, flush to keep headings clean in next loop
            if _tok_len(candidate) > MAX_TOKENS:
                _flush_text_buffer(
                    chunks,
                    current_text,
                    section_info,
                    file_hash=file_hash,
                    page_num=last_ctx["page_num"],
                    page_label=last_ctx["page_label"],
                    page_type=last_ctx["page_type"],
                    page_idx=last_ctx["page_idx"],
                    doc_name=doc_name,
                    ingested_at=ingested_at,
                    doc_id=doc_id,
                    dedup_set=doc_dedup_set,
                    drop_counters=drop_counters,
                )
                current_text["buf"] = s.strip()
            else:
                current_text["buf"] = candidate

        def _eq_num_loose(a: Optional[str], b: Optional[str]) -> bool:
            if not a or not b:
                return False
            return a.rstrip(".") == b.rstrip(".")

        def _matches_any_heading_pattern(s: str) -> bool:
            for lvl in compiled_by_level:
                for patt in compiled_by_level[lvl]:
                    if patt.match(s):
                        return True
            return False

        def _has_heading_font(spans: List[dict], thresh: float) -> bool:
            for sp in spans:
                size = sp.get("size", 0.0)
                flags = sp.get("flags", 0)
                is_bold = bool(flags & (1 << 4))
                if size > thresh or is_bold:
                    return True
            return False

        def _area(b):
            x0, y0, x1, y1 = b
            return max(0.0, (x1 - x0)) * max(0.0, (y1 - y0))

        def _iou(a, b):
            ax0, ay0, ax1, ay1 = a
            bx0, by0, bx1, by1 = b
            ix0, iy0, ix1, iy1 = max(ax0, bx0), max(ay0, by0), min(ax1, bx1), min(ay1, by1)
            iw, ih = max(0.0, ix1 - ix0), max(0.0, iy1 - iy0)
            inter = iw * ih
            if inter == 0:
                return 0.0
            return inter / (_area(a) + _area(b) - inter + 1e-9)

        def _nms(boxes, iou_thr=0.5, top_k=6):
            if not boxes:
                return []
            boxes = [(b, (_area(b) if s is None else s)) for (b, s) in boxes]
            boxes.sort(key=lambda t: t[1], reverse=True)
            kept = []
            while boxes and len(kept) < top_k:
                b, s = boxes.pop(0)
                kept.append(b)
                boxes = [(bb, ss) for (bb, ss) in boxes if _iou(b, bb) < iou_thr]
            return kept

        # ───────── Parse pages ─────────
        for page_idx in range(pdf.page_count):
            if page_idx in skip_pages:
                continue

            pg = pdf.load_page(page_idx)
            page_num = page_idx + 1

            try:
                page_label_pdf = pg.get_label() or str(page_num)
            except Exception:
                page_label_pdf = str(page_num)

            header_present, footer_present, header_band, footer_band, content_region = bands_for_page(
                pg, rules
            )
            hf_label = extract_page_numbers_from_headers_footers(
                pg, header_band, footer_band, page_numbering_info
            )
            page_label = hf_label or page_label_pdf

            is_pre_body_region = (page_idx in toc_pages) or (page_idx < body_start)

            # collect table rects for de-dup with text/images
            try:
                _tb_finder = pg.find_tables()
                _tb_list = (
                    _tb_finder.tables
                    if hasattr(_tb_finder, "tables")
                    else (_tb_finder or [])
                )
            except Exception:
                _tb_list = []
            table_rects: List[Tuple[float, float, float, float]] = []
            for _tb in _tb_list:
                try:
                    _b = _tb.bbox
                    _bb = (
                        (_b.x0, _b.y0, _b.x1, _b.y1)
                        if not isinstance(_b, tuple)
                        else _b
                    )
                    table_rects.append(_bb)
                except Exception:
                    continue

            # Decide page_type (override)
            is_toc_candidate = is_pre_body_region
            try:
                page_type = _page_kind_with_override(
                    page_idx, body_start, hf_label, is_toc_candidate
                )
            except TypeError:
                page_type = _page_kind_with_override(page_idx, body_start, hf_label)
            if (
                page_idx in toc_pages
                and hf_label
                and re.match(r"^\d+$", str(hf_label))
                and int(hf_label) >= 1
            ):
                page_type = "body"
            if page_type != "body":
                new_kind = _page_kind_with_override(page_idx, body_start, hf_label)
                page_type = new_kind

            last_ctx.update(
                {
                    "page_num": page_num,
                    "page_label": page_label,
                    "page_type": page_type,
                    "page_idx": page_idx,
                }
            )

            # ── TEXT ─────────────────────────────────────────────────────
            blocks = pg.get_text("dict").get("blocks", [])
            b = 0
            while b < len(blocks):
                blk = blocks[b]
                b += 1
                if blk.get("type") != 0 or not blk.get("bbox"):
                    continue
                x0, y0, x1, y1 = blk["bbox"]

                # skip text overlapping any table region (avoid duplicate table-as-text)
                def _overlaps(ax0, ay0, ax1, ay1, bx0, by0, bx1, by1):
                    return not (
                        ax1 <= bx0 or bx1 <= ax0 or ay1 <= by0 or by1 <= ay0
                    )

                if any(
                    _overlaps(x0, y0, x1, y1, tx0, ty0, tx1, ty1)
                    for (tx0, ty0, tx1, ty1) in table_rects
                ):
                    continue

                # clip to content region & exclude header/footer
                if x1 < content_region[0] or x0 > content_region[2]:
                    continue
                if y1 < content_region[1] or y0 > content_region[3]:
                    continue
                ym = (y0 + y1) / 2.0
                if header_present and header_band[0] <= ym <= header_band[1]:
                    continue
                if footer_present and footer_band[0] <= ym <= footer_band[1]:
                    continue

                lines = blk.get("lines", []) or []
                l = 0
                while l < len(lines):
                    ln = lines[l]
                    l += 1
                    spans = ln.get("spans", []) or []
                    line_text = clean_text(
                        "".join(sp.get("text", "") for sp in spans)
                    )
                    if not line_text:
                        continue

                    # Header/Footer de-dup (line-level)
                    if hf_filter.should_drop(line_text):
                        continue
                    hf_filter.consider(line_text, page_idx)

                    # Drop TOC-like lines
                    if is_toc_like_line(line_text):
                        dotted_drops += 1
                        continue
                    if is_pre_body_region and is_leaderless_toc_line(
                        line_text, spans=spans
                    ):
                        leaderless_drops += 1
                        continue

                    # Number+title fusion, with page-number guards
                    fused_text = line_text
                    if is_pure_number_token(line_text):
                        if hf_label and line_text.strip().rstrip(".") == str(
                            hf_label
                        ).rstrip("."):
                            continue
                        num_token = line_text.strip()
                        num_token_drop = re.sub(
                            r"[.:)\-–—]+$", "", num_token
                        ).strip()

                        def choose_fusion(num_tok: str, next_text: str) -> str:
                            cand_with_dot = f"{num_tok} {next_text}".strip()
                            cand_no_dot = f"{num_token_drop} {next_text}".strip()
                            if _matches_any_heading_pattern(cand_with_dot):
                                return cand_with_dot
                            if _matches_any_heading_pattern(cand_no_dot):
                                return cand_no_dot
                            l1_wants_dot = any(
                                p.pattern.startswith(r"^\d+\.\s+")
                                for p in compiled_by_level.get(1, [])
                            )
                            return (
                                cand_with_dot
                                if (num_tok.endswith(".") and l1_wants_dot)
                                else cand_no_dot
                            )

                        if l < len(lines):
                            nxt_ln = lines[l]
                            n_spans = nxt_ln.get("spans", []) or []
                            n_text = clean_text(
                                "".join(sp.get("text", "") for sp in n_spans)
                            )
                            if n_text:
                                fused_text = choose_fusion(num_token, n_text)
                                l += 1
                                fused_count += 1
                        elif b < len(blocks):
                            nxt_blk = blocks[b]
                            if nxt_blk.get("type") == 0 and nxt_blk.get("bbox"):
                                nbx0, nby0, nbx1, nby1 = nxt_blk["bbox"]
                                W = pg.rect.width
                                H = pg.rect.height
                                same_left = abs(nbx0 - x0) <= 0.02 * W
                                near_vert = (nby0 - y1) < 0.05 * H
                                if (
                                    same_left
                                    and near_vert
                                    and (nxt_blk.get("lines") or [])
                                ):
                                    first_ln = nxt_blk["lines"][0]
                                    n_spans = first_ln.get("spans", []) or []
                                    n_text = clean_text(
                                        "".join(
                                            sp.get("text", "") for sp in n_spans
                                        )
                                    )
                                    if n_text:
                                        fused_text = choose_fusion(num_token, n_text)
                                        del nxt_blk["lines"][0]
                                        fused_count += 1
                                        if not nxt_blk["lines"]:
                                            b += 1

                    # Heading classification
                    is_heading = False
                    detected_level = None
                    vec = None

                    for lvl in sorted(
                        compiled_by_level.keys(), reverse=True
                    ):
                        for patt in compiled_by_level[lvl]:
                            if not patt.match(fused_text):
                                continue
                            if any(
                                cre.match(fused_text) for cre in CAPTION_RE_LIST
                            ):
                                continue
                            vec = numeric_vector_from_heading(fused_text)
                            eff_depth = (
                                effective_depth_from_vector(vec) if vec else None
                            )
                            if vec and len(vec) == 1 and 1900 <= vec[0] <= 2100:
                                continue
                            detected_level = eff_depth if vec else lvl

                            # Optional whitelist gate
                            if toc_whitelist["has_toc"]:
                                cand_num = _normalize_heading_number(fused_text)
                                cand_title = _normalize_heading_title(fused_text)
                                ok_num = False
                                if cand_num and cand_num in toc_whitelist["nums"]:
                                    ok_num = True
                                elif cand_num:
                                    for n in toc_whitelist["nums"]:
                                        if _eq_num_loose(
                                            cand_num, n
                                        ) or cand_num.startswith(
                                            n.rstrip(".") + "."
                                        ):
                                            ok_num = True
                                            break
                                ok_title = bool(
                                    cand_title
                                    and cand_title in toc_whitelist["titles"]
                                )
                                if not (ok_num or ok_title):
                                    vec = None
                                    detected_level = None
                                    continue

                            is_heading = True
                            break
                        if is_heading:
                            break

                    # Unnumbered promotion via TOC (font/style guard)
                    if (
                        not is_heading
                        and toc_whitelist["has_toc"]
                        and toc_whitelist["unnumbered"]
                    ):
                        cand_title = _normalize_heading_title(fused_text)
                        for e in toc_whitelist["idx"]:
                            if cand_title == e["title_norm"]:
                                if structurally_heading_like(
                                    fused_text
                                ) and _has_heading_font(spans, heading_thresh):
                                    is_heading = True
                                    detected_level = e["level"]
                                    break

                    if is_heading:
                        # Flush current body (ensures heading starts a new clean chunk)
                        _flush_text_buffer(
                            chunks,
                            current_text,
                            section_info,
                            file_hash=file_hash,
                            page_num=page_num,
                            page_label=page_label,
                            page_type=page_type,
                            page_idx=page_idx,
                            doc_name=doc_name,
                            ingested_at=ingested_at,
                            doc_id=doc_id,
                            dedup_set=doc_dedup_set,
                            drop_counters=drop_counters,
                        )

                        # Resolve to TOC node if possible
                        cand_num = _normalize_heading_number(fused_text)
                        cand_title_norm = _normalize_heading_title(fused_text)
                        node = None
                        if cand_num and cand_num in by_num:
                            node = by_num[cand_num]
                        elif detected_level:
                            lst = by_title_at_level.get(
                                (int(detected_level), cand_title_norm)
                            ) or []
                            node = lst[0] if lst else None

                        # Reset stack to node parents
                        if node:
                            parent_len = max(0, len(node["path_titles"]) - 1)
                            for depth in range(1, parent_len + 1):
                                while (
                                    section_stack
                                    and section_stack[-1]["level"] >= depth
                                ):
                                    section_stack.pop()
                                parent_title = node["path_titles"][depth - 1]
                                parent_num = node["path_nums"][depth - 1]
                                if (
                                    section_stack
                                    and section_stack[-1]["level"] == depth
                                ):
                                    last = section_stack[-1]
                                    if (
                                        (last.get("num") or "").rstrip(".")
                                        == (parent_num or "").rstrip(".")
                                        and (
                                            last.get("title") or ""
                                        ).strip().lower()
                                        == parent_title.strip().lower()
                                    ):
                                        continue
                                section_stack.append(
                                    {
                                        "level": depth,
                                        "title": parent_title,
                                        "num": parent_num,
                                        "section_type": "synthetic",
                                    }
                                )

                        # Push current heading
                        new_level = (
                            len(node["path_titles"])
                            if node
                            else int(detected_level or 1)
                        )
                        while (
                            section_stack
                            and section_stack[-1]["level"] >= new_level
                        ):
                            section_stack.pop()

                        m_num = re.match(
                            r"^(\d+(?:\.\d+)*)\s+(.*)$", fused_text
                        )
                        if m_num:
                            num, title_only = m_num.groups()
                        else:
                            num, title_only = None, fused_text
                        if node:
                            num = node.get("num") or num

                        section_stack.append(
                            {
                                "level": new_level,
                                "title": title_only,
                                "num": num,
                                "section_type": "content",
                            }
                        )

                    else:
                        # Body text
                        append_body_text(fused_text)

            # ── TABLES (structured) ──────────────────────────────────────
            try:
                tables = pg.find_tables()
                table_list = (
                    tables.tables
                    if hasattr(tables, "tables")
                    else (tables or [])
                )
            except Exception:
                table_list = []
            for tb in table_list:
                try:
                    matrix = tb.extract()
                    if not matrix:
                        continue
                    header = matrix[0] if matrix else []
                    rows = matrix[1:] if len(matrix) > 1 else []
                    tb_bbox = None
                    if hasattr(tb, "bbox"):
                        bbb = tb.bbox
                        tb_bbox = (
                            (bbb.x0, bbb.y0, bbb.x1, bbb.y1)
                            if not isinstance(bbb, tuple)
                            else bbb
                        )
                    caption = (
                        find_pdf_caption_near(pg, tb_bbox) if tb_bbox else None
                    )
                    si = section_info()

                    # Deterministic ID for tables too
                    cid = (
                        f"{doc_id}::{page_label}::t0::tbl"
                        f"{hashlib.sha1(str(tb_bbox).encode()).hexdigest()[:6]}"
                    )

                    chunks.append(
                        {
                            **si,
                            "doc_name": doc_name,
                            "doc_type": "pdf",
                            "file_hash": file_hash,
                            "doc_id": doc_id,
                            "page_index": page_num,
                            "page_label": page_label,
                            "page_type": page_type,
                            "modality": "table",
                            "kind": "table",
                            "section_title": si.get("section"),
                            "heading_path": si.get("section_path"),
                            "content": json.dumps(
                                {"columns": header, "rows": rows},
                                ensure_ascii=False,
                            ),
                            "caption_text": caption,
                            "bounding_box": list(tb_bbox) if tb_bbox else None,
                            "ingested_at": ingested_at,
                            "chunk_id": cid,
                            "chunk_hash": cid,
                        }
                    )

                    if section_stack and re.match(
                        r"^\s*Table\s+\d+\b",
                        section_stack[-1].get("title", ""),
                        re.I,
                    ):
                        section_stack.pop()

                except Exception as e:
                    dbg(f"[DBG] Table extraction error p{page_idx}: {e}")

            # ── VISUALS (images / charts via GPT image JSON, old behavior) ─────
            img_candidates: List[Tuple[Tuple[float, float, float, float], Optional[float]]] = []
            seen_xref_boxes = set()
            for xref, *_ in pg.get_images(full=True):
                try:
                    for rect in pg.get_image_rects(xref) or []:
                        bbox = (rect.x0, rect.y0, rect.x1, rect.y1)
                        # skip full-page
                        if (
                            abs(bbox[0] - pg.rect.x0) < 1
                            and abs(bbox[1] - pg.rect.y0) < 1
                            and abs(bbox[2] - pg.rect.x1) < 1
                            and abs(bbox[3] - pg.rect.y1) < 1
                        ):
                            continue
                        # skip if overlaps table (avoid duplicating charts inside table regions)
                        if any(
                            _iou(bbox, tr) > 0.2 for tr in table_rects
                        ):
                            continue
                        w, h = (rect.x1 - rect.x0), (rect.y1 - rect.y0)
                        if w < 30 or h < 30:
                            continue
                        area_ratio = (w * h) / (
                            pg.rect.width * pg.rect.height
                        )
                        if area_ratio < 0.01:
                            continue
                        key = (
                            round(bbox[0], 1),
                            round(bbox[1], 1),
                            round(bbox[2], 1),
                            round(bbox[3], 1),
                        )
                        if key in seen_xref_boxes:
                            continue
                        seen_xref_boxes.add(key)
                        img_candidates.append((bbox, None))
                except Exception as e:
                    dbg(f"[DBG] Image rects error p{page_idx}: {e}")

            kept_img_boxes = _nms(img_candidates, iou_thr=0.5, top_k=6)

            draw_candidates: List[Tuple[Tuple[float, float, float, float], Optional[float]]] = []
            try:
                drawings = pg.get_drawings() or []
            except Exception:
                drawings = []
            for dr in drawings:
                try:
                    r = dr.get("rect", None)
                    if not r:
                        continue
                    bbox = (r.x0, r.y0, r.x1, r.y1)
                    # skip full-page
                    if (
                        abs(bbox[0] - pg.rect.x0) < 1
                        and abs(bbox[1] - pg.rect.y0) < 1
                        and abs(bbox[2] - pg.rect.x1) < 1
                        and abs(bbox[3] - pg.rect.y1) < 1
                    ):
                        continue
                    if any(_iou(bbox, tr) > 0.2 for tr in table_rects):
                        continue
                    w, h = (r.x1 - r.x0), (r.y1 - r.y0)
                    if w < 40 or h < 40:
                        continue
                    area_ratio = (w * h) / (
                        pg.rect.width * pg.rect.height
                    )
                    if area_ratio < 0.02:
                        continue
                    draw_candidates.append((bbox, None))
                except Exception as e:
                    dbg(f"[DBG] Drawing rect error p{page_idx}: {e}")

            kept_draw_boxes = _nms(draw_candidates, iou_thr=0.5, top_k=4)
            filtered_draw_boxes = [
                db
                for db in kept_draw_boxes
                if all(_iou(db, ib) < 0.3 for ib in kept_img_boxes)
            ]

            def _emit_visual(pg: fitz.Page, bbox: Tuple[float, float, float, float]):
                # Old behavior: GPT JSON → code/table/image
                crop_img = render_bbox_image(pg, bbox, dpi=192)  # -> PIL.Image or None
                vision = None
                if crop_img is not None and RUN_GPT_IMAGE_DESCRIBE and client is not None:
                    try:
                        buf = io.BytesIO()
                        crop_img.save(buf, format="PNG")
                        blob = buf.getvalue()
                        vision = _describe_image_with_gpt(
                            blob, mime_hint="image/png"
                        )
                        dbg(
                            f"[DBG] PDF GPT-Vision called: {bool(vision)} | keys={list(vision.keys()) if vision else None}"
                        )
                    except Exception as e:
                        dbg(f"[DBG] _describe_image_with_gpt failed: {e}")
                        import traceback

                        traceback.print_exc()

                # Caption: prefer GPT caption, else nearby PDF caption, else placeholder
                caption = None
                if vision and vision.get("caption"):
                    caption = vision.get("caption")
                else:
                    caption = (
                        find_pdf_caption_near(pg, bbox)
                        or "[No caption found — auto placeholder]"
                    )
                dbg(
                    f"[DBG] PDF Caption resolved: {caption[:80]}..."
                    if caption
                    else "[DBG] PDF Caption missing"
                )

                si = section_info()

                # Upgrade: code from image
                if (
                    vision
                    and vision.get("modality") == "code"
                    and (vision.get("code") or "").strip()
                ):
                    content_text = vision.get("code", "")
                    img_sig = hashlib.sha1(
                        (str(bbox) + content_text).encode()
                    ).hexdigest()[:8]
                    cid = f"{doc_id}::{page_label}::img:{img_sig}"
                    chunks.append(
                        {
                            **si,
                            "doc_name": doc_name,
                            "doc_type": "pdf",
                            "file_hash": file_hash,
                            "doc_id": doc_id,
                            "page_index": page_num,
                            "page_label": page_label,
                            "page_type": page_type,
                            "modality": "code",
                            "kind": "code",
                            "section_title": si.get("section"),
                            "heading_path": si.get("section_path"),
                            "content": content_text,
                            "caption_text": caption,
                            "bounding_box": list(bbox),
                            "ingested_at": ingested_at,
                            "chunk_id": cid,
                            "chunk_hash": cid,
                        }
                    )
                    return

                # Upgrade: table from image
                if (
                    vision
                    and vision.get("modality") == "table"
                    and ((vision.get("columns") or []) or (vision.get("rows") or []))
                ):
                    content_obj = {
                        "columns": vision.get("columns") or [],
                        "rows": vision.get("rows") or [],
                    }
                    content_json = json.dumps(
                        content_obj, ensure_ascii=False
                    )
                    img_sig = hashlib.sha1(
                        (str(bbox) + content_json).encode()
                    ).hexdigest()[:8]
                    cid = f"{doc_id}::{page_label}::imgtbl:{img_sig}"
                    chunks.append(
                        {
                            **si,
                            "doc_name": doc_name,
                            "doc_type": "pdf",
                            "file_hash": file_hash,
                            "doc_id": doc_id,
                            "page_index": page_num,
                            "page_label": page_label,
                            "page_type": page_type,
                            "modality": "table",
                            "kind": "table",
                            "section_title": si.get("section"),
                            "heading_path": si.get("section_path"),
                            "content": content_json,
                            "caption_text": caption,
                            "bounding_box": list(bbox),
                            "ingested_at": ingested_at,
                            "chunk_id": cid,
                            "chunk_hash": cid,
                        }
                    )
                    # Avoid scoping following prose under a "Table N ..." section title
                    if section_stack and re.match(
                        r"^\s*Table\s+\d+\b",
                        section_stack[-1].get("title", ""),
                        re.I,
                    ):
                        section_stack.pop()
                    return

                # Plain image/chart/diagram → long description in content
                long_desc = ""
                if vision and vision.get("long_description"):
                    long_desc = vision.get("long_description")
                else:
                    long_desc = caption or ""
                if long_desc and len(long_desc) > 8000:
                    long_desc = long_desc[:8000]

                img_sig = hashlib.sha1(
                    (str(bbox) + long_desc).encode()
                ).hexdigest()[:8]
                cid = f"{doc_id}::{page_label}::imgdesc:{img_sig}"
                chunks.append(
                    {
                        **si,
                        "doc_name": doc_name,
                        "doc_type": "pdf",
                        "file_hash": file_hash,
                        "doc_id": doc_id,
                        "page_index": page_num,
                        "page_label": page_label,
                        "page_type": page_type,
                        "modality": "image",
                        "kind": "image",
                        "section_title": si.get("section"),
                        "heading_path": si.get("section_path"),
                        "content": long_desc,
                        "caption_text": caption,
                        "bounding_box": list(bbox),
                        "ingested_at": ingested_at,
                        "chunk_id": cid,
                        "chunk_hash": cid,
                    }
                )
                # If the last section is literally "Table N ...", don't let it scope prose
                if section_stack and re.match(
                    r"^\s*Table\s+\d+\b",
                    section_stack[-1].get("title", ""),
                    re.I,
                ):
                    section_stack.pop()

            for bbox in kept_img_boxes:
                _emit_visual(pg, bbox)
            for bbox in filtered_draw_boxes:
                _emit_visual(pg, bbox)

        # End-of-document flush
        if (current_text["buf"] or "").strip():
            _flush_text_buffer(
                chunks,
                current_text,
                section_info,
                file_hash=file_hash,
                page_num=last_ctx["page_num"],
                page_label=last_ctx["page_label"],
                page_type=last_ctx["page_type"],
                page_idx=last_ctx["page_idx"],
                doc_name=doc_name,
                ingested_at=ingested_at,
                doc_id=doc_id,
                dedup_set=doc_dedup_set,
                drop_counters=drop_counters,
            )

        if DEBUG:
            dbg(
                f"[DBG] Repeating header/footer lines dropped: {hf_filter.dropped_count}"
            )
            dbg(f"[DBG] Quality drops: {drop_counters}")

        return chunks
    finally:
        pdf.close()


# ────────────────────────── DOCX parsing (IDs/meta aligned with PDF) ─────────
def parse_docx_to_chunks(path: str) -> List[dict]:
    from docx import Document as DocxDoc
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.ns import qn

    def int_to_roman(n: int, upper=True) -> str:
        vals = [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]
        out = []
        x = max(1, int(n))
        for v, s in vals:
            while x >= v:
                out.append(s)
                x -= v
        s = "".join(out)
        return s if upper else s.lower()

    def int_to_alpha(n: int, upper=True) -> str:
        n = max(1, int(n))
        letters = []
        while n > 0:
            n, r = divmod(n - 1, 26)
            letters.append(chr((65 if upper else 97) + r))
        return "".join(reversed(letters))

    def fmt_label(n: int, fmt: str) -> str:
        f = (fmt or "decimal").lower()
        if f == "decimal":
            return str(n)
        if f in ("upperroman", "roman"):
            return int_to_roman(n, upper=True)
        if f == "lowerroman":
            return int_to_roman(n, upper=False)
        if f == "upperletter":
            return int_to_alpha(n, upper=True)
        if f == "lowerletter":
            return int_to_alpha(n, upper=False)
        return str(n)

    path = Path(path)
    doc = DocxDoc(str(path))
    file_hash = sha256_file(path)
    doc_name = path.name
    doc_id = _make_doc_id(path, file_hash)

    chunks: List[dict] = []
    section_stack: List[dict] = []
    current_text = {"buf": ""}  # accumulate body text within page
    ingested_at = now_utc_iso()
    doc_dedup_set: set = set()
    drop_counters = {"too_short_chars": 0, "too_short_tokens": 0, "exact_dup": 0}

    page_num: int = 1
    page_fmt: str = "decimal"
    last_ctx = {"page_index": page_num, "page_label": fmt_label(page_num, page_fmt)}

    def section_info():
        if not section_stack:
            return {
                "section": "Document",
                "section_level": 0,
                "section_path": "Document",
                "section_type": "content",
            }
        path_parts: List[str] = []
        acc: List[dict] = []
        for s in section_stack:
            while acc and acc[-1]["level"] >= s["level"]:
                acc.pop()
                if path_parts:
                    path_parts.pop()
            acc.append(s)
            title_clean = re.sub(
                r"^\d+(?:\.\d+)*\s+", "", s["title"]
            ).strip()
            path_parts.append(title_clean or s["title"])
        return {
            "section": section_stack[-1]["title"],
            "section_level": section_stack[-1]["level"],
            "section_path": " / ".join(p for p in path_parts if p),
            "section_type": section_stack[-1].get("section_type", "content"),
        }

    def _emit_text_chunk(text: str, pidx: int, plabel: str):
        if not (text or "").strip():
            return
        modality = "code" if _is_probably_code(text) else "text"
        # deterministic windows + IDs
        for piece, start_off in iter_token_windows(text, MAX_TOKENS, OVERLAP):
            if len(piece) < MIN_CHARS_PER_CHUNK:
                drop_counters["too_short_chars"] += 1
                continue
            if _tok_len(piece) < MIN_TOKENS_PER_CHUNK:
                drop_counters["too_short_tokens"] += 1
                continue
            ph = hashlib.sha1(piece.encode("utf-8")).hexdigest()
            if ph in doc_dedup_set:
                drop_counters["exact_dup"] += 1
                continue
            doc_dedup_set.add(ph)

            ph_short = ph[:8]
            cid = f"{doc_id}::{plabel}::t{start_off}:{ph_short}"
            si = section_info()
            chunks.append(
                {
                    **si,
                    "doc_name": doc_name,
                    "doc_type": "docx",
                    "file_hash": file_hash,
                    "doc_id": doc_id,
                    "page_index": pidx,
                    "page_label": plabel,
                    "page_type": "body",
                    "modality": modality,
                    "kind": ("code" if modality == "code" else "text"),
                    "section_title": si.get("section"),
                    "heading_path": si.get("section_path"),
                    "content": piece,
                    "caption_text": "None",
                    "bounding_box": "None",
                    "ingested_at": ingested_at,
                    "chunk_id": cid,
                    "chunk_hash": cid,
                }
            )

    def flush_text():
        buf = current_text["buf"]
        if buf.strip():
            _emit_text_chunk(
                buf, last_ctx["page_index"], last_ctx["page_label"]
            )
        current_text["buf"] = ""

    def _low_signal_line(s: str) -> bool:
        s2 = (s or "").strip()
        if not s2:
            return True
        if len(s2) <= 2:
            return True
        return False

    def _is_heading_para(para: DocxParagraph) -> Tuple[bool, Optional[int]]:
        text = clean_text(para.text)
        if not text:
            return False, None
        style_name = (para.style.name if para.style else "") or ""
        m_style = re.match(r"Heading\s+(\d+)", style_name, re.I)
        if m_style:
            return True, int(m_style.group(1))
        m_num = re.match(r"^(\d+(?:\.\d+)*)\s+", text)
        if m_num:
            try:
                vec = [int(x) for x in m_num.group(1).split(".") if x != ""]
            except Exception:
                vec = None
            lvl = effective_depth_from_vector(vec) if vec else 1
            return True, int(lvl or 1)
        return False, None

    def _norm_title(s: str) -> str:
        return re.sub(
            r"^\d+(?:\.\d+)*\s+", "", clean_text(s or "")
        ).strip().lower()

    def _para_images(para: DocxParagraph) -> List[Tuple[bytes, Optional[str]]]:
        out: List[Tuple[bytes, Optional[str]]] = []
        for run in para.runs:
            blips = run._r.xpath(".//a:blip")
            for bl in blips:
                rId = bl.get(qn("r:embed"))
                if not rId:
                    continue
                try:
                    part = para.part.related_parts.get(rId)
                    if part and getattr(part, "blob", None):
                        ctype = getattr(part, "content_type", None)
                        out.append((bytes(part.blob), ctype))
                except Exception:
                    continue
        return out

    def _cell_text(cell) -> str:
        parts = [clean_text(p.text) for p in cell.paragraphs]
        return " ".join([p for p in parts if p])

    def _emit_table(tbl: DocxTable):
        flush_text()
        try:
            rows = []
            for r in tbl.rows:
                rows.append([_cell_text(c) for c in r.cells])

            if not rows:
                return
            columns: List[str] = []
            data_rows: List[List[str]] = []
            first_non_empty = all((cell or "").strip() for cell in rows[0])
            if first_non_empty:
                columns = [
                    c if c else f"col{i+1}"
                    for i, c in enumerate(rows[0])
                ]
                data_rows = rows[1:]
            else:
                ncols = max((len(r) for r in rows), default=0)
                columns = [f"col{i+1}" for i in range(ncols)]
                data_rows = rows

            si = section_info()
            cid = (
                f"{doc_id}::{last_ctx['page_label']}::t0::tbl"
                f"{hashlib.sha1(str(columns).encode()).hexdigest()[:6]}"
            )
            chunks.append(
                {
                    **si,
                    "doc_name": doc_name,
                    "doc_type": "docx",
                    "file_hash": file_hash,
                    "doc_id": doc_id,
                    "page_index": last_ctx["page_index"],
                    "page_label": last_ctx["page_label"],
                    "page_type": "body",
                    "modality": "table",
                    "kind": "table",
                    "section_title": si.get("section"),
                    "heading_path": si.get("section_path"),
                    "content": json.dumps(
                        {"columns": columns, "rows": data_rows},
                        ensure_ascii=False,
                    ),
                    "caption_text": "None",
                    "bounding_box": "None",
                    "ingested_at": ingested_at,
                    "chunk_id": cid,
                    "chunk_hash": cid,
                }
            )
        except Exception:
            return

    def is_caption_line(s: str) -> bool:
        if not s:
            return False
        for cre in CAPTION_RE_LIST:
            if cre.match(s):
                return True
        return False

    def _emit_image(
        blob: bytes,
        caption_hint: Optional[str] = None,
        mime_hint: Optional[str] = None,
    ):
        flush_text()
        si = section_info()
        # Old DOCX GPT-Vision behavior
        vision = _describe_image_with_gpt(blob, mime_hint=mime_hint)
        dbg(
            f"[DBG] DOCX GPT-Vision called: {bool(vision)} | keys={list(vision.keys()) if vision else None}"
        )

        # code-from-image
        if (
            vision
            and vision.get("modality") == "code"
            and (vision.get("code") or "").strip()
        ):
            content_text = vision.get("code", "")
            sig = hashlib.sha1(
                (content_text or "").encode()
            ).hexdigest()[:8]
            cid = f"{doc_id}::{last_ctx['page_label']}::img:{sig}"
            chunks.append(
                {
                    **si,
                    "doc_name": doc_name,
                    "doc_type": "docx",
                    "file_hash": file_hash,
                    "doc_id": doc_id,
                    "page_index": last_ctx["page_index"],
                    "page_label": last_ctx["page_label"],
                    "page_type": "body",
                    "modality": "code",
                    "kind": "code",
                    "section_title": si.get("section"),
                    "heading_path": si.get("section_path"),
                    "content": content_text,
                    "caption_text": vision.get("caption") or caption_hint,
                    "bounding_box": None,
                    "ingested_at": ingested_at,
                    "chunk_id": cid,
                    "chunk_hash": cid,
                }
            )
            return

        # table-from-image
        if (
            vision
            and vision.get("modality") == "table"
            and ((vision.get("columns") or []) or (vision.get("rows") or []))
        ):
            content_obj = {
                "columns": vision.get("columns") or [],
                "rows": vision.get("rows") or [],
            }
            sig = hashlib.sha1(
                json.dumps(content_obj, ensure_ascii=False).encode()
            ).hexdigest()[:8]
            cid = f"{doc_id}::{last_ctx['page_label']}::imgtbl:{sig}"
            chunks.append(
                {
                    **si,
                    "doc_name": doc_name,
                    "doc_type": "docx",
                    "file_hash": file_hash,
                    "doc_id": doc_id,
                    "page_index": last_ctx["page_index"],
                    "page_label": last_ctx["page_label"],
                    "page_type": "body",
                    "modality": "table",
                    "kind": "table",
                    "section_title": si.get("section"),
                    "heading_path": si.get("section_path"),
                    "content": json.dumps(
                        content_obj, ensure_ascii=False
                    ),
                    "caption_text": vision.get("caption") or caption_hint,
                    "bounding_box": "None",
                    "ingested_at": ingested_at,
                    "chunk_id": cid,
                    "chunk_hash": cid,
                }
            )
            return

        # plain image → long description into content, short caption in caption_text
        long_desc = ""
        if vision and vision.get("long_description"):
            long_desc = vision.get("long_description")
        else:
            long_desc = (vision and vision.get("caption")) or (caption_hint or "")
        if len(long_desc) > 8000:
            long_desc = long_desc[:8000]

        sig = hashlib.sha1(
            (long_desc or "").encode()
        ).hexdigest()[:8]
        cid = f"{doc_id}::{last_ctx['page_label']}::imgdesc:{sig}"
        chunks.append(
            {
                **si,
                "doc_name": doc_name,
                "doc_type": "docx",
                "file_hash": file_hash,
                "doc_id": doc_id,
                "page_index": last_ctx["page_index"],
                "page_label": last_ctx["page_label"],
                "page_type": "body",
                "modality": "image",
                "kind": "image",
                "section_title": si.get("section"),
                "heading_path": si.get("section_path"),
                "content": long_desc,
                "caption_text": (vision and vision.get("caption")) or caption_hint,
                "bounding_box": "None",
                "ingested_at": ingested_at,
                "chunk_id": cid,
                "chunk_hash": cid,
            }
        )

    def _para_has_page_breaks(p: DocxParagraph) -> int:
        cnt = 0
        cnt += len(p._p.xpath(".//w:br[@w:type='page']"))
        cnt += len(p._p.xpath(".//w:lastRenderedPageBreak"))
        return cnt

    def _maybe_update_section_numbering_from_para(p: DocxParagraph):
        nonlocal page_num, page_fmt, last_ctx
        sectPr = p._p.xpath("./w:pPr/w:sectPr")
        if not sectPr:
            return
        sp = sectPr[0]
        pgNumType = sp.xpath("./w:pgNumType")
        if pgNumType:
            node = pgNumType[0]
            start = node.get(qn("w:start"))
            fmt = node.get(qn("w:fmt"))
            if fmt:
                page_fmt = fmt
            if start:
                try:
                    page_num = int(start)
                except Exception:
                    page_num = 1
            last_ctx["page_index"] = page_num
            last_ctx["page_label"] = fmt_label(page_num, page_fmt)

    def iter_block_items(parent):
        parent_elm = parent.element if hasattr(parent, "element") else parent._element
        for child in parent_elm.body.iterchildren():
            if isinstance(child, CT_P):
                yield DocxParagraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield DocxTable(child, parent)

    previous_para_text: Optional[str] = None

    for blk in iter_block_items(doc):
        if isinstance(blk, DocxParagraph):
            _maybe_update_section_numbering_from_para(blk)

        if isinstance(blk, DocxParagraph):
            breaks_here = _para_has_page_breaks(blk)
            if breaks_here > 0:
                flush_text()
                page_num += breaks_here
                last_ctx["page_index"] = page_num
                last_ctx["page_label"] = fmt_label(page_num, page_fmt)

        if isinstance(blk, DocxTable):
            for row in blk.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        blobs = _para_images(p)
                        if blobs:
                            cap_here = clean_text(p.text)
                            cap_near = (
                                cap_here if is_caption_line(cap_here) else None
                            )
                            for blob, ctype in blobs:
                                _emit_image(
                                    blob,
                                    caption_hint=cap_near
                                    or (
                                        previous_para_text
                                        if is_caption_line(
                                            previous_para_text or ""
                                        )
                                        else None
                                    ),
                                    mime_hint=ctype,
                                )
            _emit_table(blk)
            previous_para_text = None
            continue

        if isinstance(blk, DocxParagraph):
            text = clean_text(blk.text)

            if is_toc_like_line(text):
                previous_para_text = text
                continue

            blobs = _para_images(blk)
            if blobs:
                cap_here = text if is_caption_line(text) else None
                cap_prev = (
                    previous_para_text
                    if is_caption_line(previous_para_text or "")
                    else None
                )
                for blob, ctype in blobs:
                    _emit_image(
                        blob,
                        caption_hint=cap_here or cap_prev,
                        mime_hint=ctype,
                    )
                if not text or is_caption_line(text) or len(text) <= 2:
                    previous_para_text = text
                    continue

            is_heading, lvl = _is_heading_para(blk)
            if is_heading and lvl:
                title_norm = _norm_title(text)
                top_same = (
                    section_stack
                    and section_stack[-1]["level"] == int(lvl)
                    and _norm_title(section_stack[-1]["title"]) == title_norm
                )
                if not top_same:
                    flush_text()
                    while (
                        section_stack
                        and section_stack[-1]["level"] >= int(lvl)
                    ):
                        section_stack.pop()
                    section_stack.append(
                        {
                            "level": int(lvl),
                            "title": text,
                            "section_type": "content",
                        }
                    )
                previous_para_text = text
                continue

            if not text or len(text) <= 2:
                previous_para_text = text
                continue

            buf = current_text["buf"]
            current_text["buf"] = (buf + "\n" + text).strip() if buf else text
            previous_para_text = text

    flush_text()

    dbg(
        f"[DBG] DOCX parsed: {doc_name} -> {len(chunks)} chunks "
        f"(sections={sorted(set(c.get('section_level',0) for c in chunks))}, "
        f"last_page={last_ctx['page_index']} {last_ctx['page_label']})"
    )

    if DEBUG:
        dbg(f"[DBG] DOCX Quality drops: {drop_counters}")

    return chunks


# ────────────────────────── Text aggregation for summaries ────────────────────
def _collect_doc_text_pieces(chunks: List[dict], doc_name: str) -> List[str]:
    pieces: List[str] = []
    for ch in chunks:
        if ch.get("doc_name") != doc_name:
            continue
        mod = ch.get("modality")
        if mod == "table":
            payload = ch.get("content")
            if isinstance(payload, str):
                try:
                    payload = json.loads(payload)
                except Exception:
                    payload = {}
            cols = payload.get("columns") or []
            rows = (payload.get("rows") or [])[:8]
            head = " | ".join(str(c) for c in cols)
            body = "\n".join(" | ".join(str(c) for c in r) for r in rows)
            block = f"Table:\n{head}\n{body}".strip()
            if block:
                pieces.append(block)
        elif mod in ("text", "code"):
            txt = ch.get("content", "")
            if txt:
                pieces.append(("Code:\n" + txt) if mod == "code" else txt)
        elif mod in ("image", "chart", "image_desc"):
            cap = ch.get("caption_text")
            if cap:
                pieces.append(f"Figure: {cap}")
    return pieces


# ────────────────────────── Chroma settings, mapping ──────────────────────────
client_settings = Settings(
    is_persistent=True,
    persist_directory=STORE_DIR,
)


def _chunk_to_document(ch: Dict) -> Document:
    content = (ch.get("content") or "").strip()
    if not content:
        return Document(page_content="", metadata={})

    meta = {
        k: (
            "None"
            if v is None
            else (
                json.dumps(v, ensure_ascii=False)
                if isinstance(v, (list, dict, tuple, set))
                else v
            )
        )
        for k, v in ch.items()
        if k != "content"
    }

    doc_name = meta.get("doc_name", "")
    default_source = (
        os.path.splitext(doc_name)[0]
        if doc_name
        else meta.get("file_hash", "")[:10]
    )
    meta["source"] = meta.get("source", default_source)

    if "doc_id" not in meta:
        fh = meta.get("file_hash", "")
        meta["doc_id"] = (
            f"{meta['source']}::{fh[:10]}" if fh else meta["source"]
        )

    return Document(page_content=content, metadata=meta)


def to_documents_and_ids(chunks: List[Dict]) -> TypingTuple[List[Document], List[str]]:
    docs: List[Document] = []
    ids: List[str] = []

    for ch in chunks:
        d = _chunk_to_document(ch)
        if not (d.page_content or "").strip():
            continue

        m = d.metadata or {}
        cid = m.get("chunk_id") or m.get("chunk_hash")
        if not cid:
            base = f"{m.get('source','')}|{m.get('doc_name','')}|{d.page_content[:120]}"
            cid = hashlib.sha1(base.encode("utf-8")).hexdigest()

        docs.append(d)
        ids.append(str(cid))

    return docs, ids
